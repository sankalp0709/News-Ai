import io
import logging
from typing import Dict, Any
import docx

logger = logging.getLogger(__name__)

class DocumentDocxAdapter:
    """DOCX extraction adapter."""
    
    @staticmethod
    def extract(file_content: bytes, filename: str) -> Dict[str, Any]:
        if not file_content:
            raise ValueError("Empty DOCX file")
            
        try:
            doc = docx.Document(io.BytesIO(file_content))
            
            # Extract paragraphs
            text_chunks = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_chunks.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_data:
                        text_chunks.append(" | ".join(row_data))
                        
            text = "\n".join(text_chunks)
            
            # Extract basic metadata
            metadata = {}
            if doc.core_properties:
                props = doc.core_properties
                if getattr(props, "title", None): metadata["title"] = props.title
                if getattr(props, "author", None): metadata["author"] = props.author
                if getattr(props, "subject", None): metadata["subject"] = props.subject
                if getattr(props, "created", None): metadata["created"] = str(props.created)
                if getattr(props, "modified", None): metadata["modified"] = str(props.modified)
                    
            return {
                "format": "docx",
                "filename": filename,
                "content": {
                    "text": text.strip(),
                    "metadata": metadata
                }
            }
        except ValueError:
            raise
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
